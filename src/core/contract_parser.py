"""
NEEX Legal Contract Review System
Document Parser Module

Handles parsing of legal contracts from various formats (PDF, DOCX, TXT)
Extracts structured clause information for analysis.
"""

import re
import logging
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None

logger = logging.getLogger(__name__)


class DocumentFormat(Enum):
    """Supported document formats."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"


@dataclass
class Clause:
    """Represents a contract clause with metadata."""
    number: int
    title: str
    content: str
    section: Optional[str] = None
    page_number: Optional[int] = None
    char_position: Tuple[int, int] = (0, 0)  # start, end positions
    
    def __post_init__(self):
        """Clean and validate clause content."""
        self.content = self.content.strip()
        self.title = self.title.strip()


@dataclass  
class ContractDocument:
    """Represents a parsed contract document."""
    title: str
    clauses: List[Clause]
    metadata: Dict[str, str]
    source_file: Path
    format: DocumentFormat
    
    @property
    def total_clauses(self) -> int:
        """Return total number of clauses."""
        return len(self.clauses)
    
    @property
    def total_content_length(self) -> int:
        """Return total character count of all clauses."""
        return sum(len(clause.content) for clause in self.clauses)


class ContractParser:
    """Main parser class for legal contract documents."""
    
    def __init__(self):
        """Initialize parser with default clause patterns."""
        self.clause_patterns = [
            r'^(\d+)\.\s+([^\n]+)',  # 1. Title
            r'^Section\s+(\d+):?\s*([^\n]*)',  # Section 1: Title
            r'^Article\s+(\d+):?\s*([^\n]*)',  # Article 1: Title
            r'^([a-zA-Z])\.\s+([^\n]+)',  # a. Title
            r'^\(([a-zA-Z0-9]+)\)\s+([^\n]*)',  # (a) Title
        ]
        
    def parse_document(self, file_path: Path) -> ContractDocument:
        """
        Parse a contract document from file.
        
        Args:
            file_path: Path to the contract file
            
        Returns:
            ContractDocument with parsed clauses
            
        Raises:
            ValueError: If format not supported
            FileNotFoundError: If file doesn't exist
        """
        if not file_path.exists():
            raise FileNotFoundError(f"Contract file not found: {file_path}")
            
        format_type = self._detect_format(file_path)
        logger.info(f"Parsing {format_type.value} document: {file_path.name}")
        
        # Extract raw text based on format
        if format_type == DocumentFormat.PDF:
            text = self._parse_pdf(file_path)
        elif format_type == DocumentFormat.DOCX:
            text = self._parse_docx(file_path)
        elif format_type == DocumentFormat.TXT:
            text = self._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported format: {format_type}")
            
        # Extract clauses from text
        clauses = self._extract_clauses(text)
        
        # Extract metadata
        metadata = self._extract_metadata(text, file_path)
        
        return ContractDocument(
            title=metadata.get('title', file_path.stem),
            clauses=clauses,
            metadata=metadata,
            source_file=file_path,
            format=format_type
        )
    
    def _detect_format(self, file_path: Path) -> DocumentFormat:
        """Detect document format from file extension."""
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            if not PyPDF2:
                raise ImportError("PyPDF2 required for PDF parsing. Install with: pip install PyPDF2")
            return DocumentFormat.PDF
        elif suffix in ['.docx', '.doc']:
            if not Document:
                raise ImportError("python-docx required for DOCX parsing. Install with: pip install python-docx")
            return DocumentFormat.DOCX
        elif suffix == '.txt':
            return DocumentFormat.TXT
        else:
            raise ValueError(f"Unsupported file format: {suffix}")
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Extract text from PDF file."""
        text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    text += f"\\n[PAGE {page_num + 1}]\\n{page_text}"
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
            
        return text
    
    def _parse_docx(self, file_path: Path) -> str:
        """Extract text from DOCX file."""
        try:
            doc = Document(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text.strip())
                    
            return "\\n".join(paragraphs)
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
    
    def _parse_txt(self, file_path: Path) -> str:
        """Read text from TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin1') as file:
                return file.read()
    
    def _extract_clauses(self, text: str) -> List[Clause]:
        """
        Extract clauses from contract text using pattern matching.
        
        Args:
            text: Full contract text
            
        Returns:
            List of identified clauses
        """
        clauses = []
        lines = text.split('\\n')
        current_clause = None
        clause_number = 0
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if line starts a new clause
            clause_match = self._match_clause_header(line)
            
            if clause_match:
                # Save previous clause if exists
                if current_clause and current_clause['content'].strip():
                    clause_number += 1
                    clauses.append(Clause(
                        number=clause_number,
                        title=current_clause['title'],
                        content=current_clause['content'].strip(),
                        section=current_clause.get('section'),
                        char_position=(current_clause['start_pos'], i)
                    ))
                
                # Start new clause
                current_clause = {
                    'title': clause_match['title'],
                    'content': '',
                    'section': clause_match.get('section'),
                    'start_pos': i
                }
            else:
                # Add to current clause content
                if current_clause:
                    current_clause['content'] += line + "\\n"
        
        # Add final clause
        if current_clause and current_clause['content'].strip():
            clause_number += 1
            clauses.append(Clause(
                number=clause_number,
                title=current_clause['title'],
                content=current_clause['content'].strip(),
                section=current_clause.get('section'),
                char_position=(current_clause['start_pos'], len(lines))
            ))
        
        logger.info(f"Extracted {len(clauses)} clauses from document")
        return clauses
    
    def _match_clause_header(self, line: str) -> Optional[Dict[str, str]]:
        """
        Check if line matches a clause header pattern.
        
        Args:
            line: Text line to check
            
        Returns:
            Dict with clause info if match found, None otherwise
        """
        for pattern in self.clause_patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                groups = match.groups()
                
                if len(groups) >= 2:
                    return {
                        'number': groups[0],
                        'title': groups[1] if groups[1] else f"Clause {groups[0]}",
                        'section': groups[0] if 'Section' in pattern or 'Article' in pattern else None
                    }
                else:
                    return {
                        'number': groups[0] if groups else '1',
                        'title': line,
                        'section': None
                    }
        return None
    
    def _extract_metadata(self, text: str, file_path: Path) -> Dict[str, str]:
        """
        Extract metadata from contract text.
        
        Args:
            text: Full contract text
            file_path: Source file path
            
        Returns:
            Dictionary of metadata
        """
        metadata = {
            'source_file': str(file_path),
            'file_size': str(file_path.stat().st_size),
            'character_count': str(len(text))
        }
        
        # Try to extract title from first few lines
        lines = text.split('\\n')[:10]
        for line in lines:
            line = line.strip()
            if len(line) > 10 and len(line) < 200:
                # Likely a title
                if any(word in line.upper() for word in ['AGREEMENT', 'CONTRACT', 'SERVICE']):
                    metadata['title'] = line
                    break
        
        # Extract parties (basic pattern matching)
        party_patterns = [
            r'between\\s+([^\\n]+?)\\s+and\\s+([^\\n]+?)(?:\\.|,|\\n)',
            r'Party\\s+A[:\s]+([^\\n]+)',
            r'Party\\s+B[:\s]+([^\\n]+)'
        ]
        
        for pattern in party_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            if matches:
                if isinstance(matches[0], tuple):
                    metadata['party_1'] = matches[0][0].strip()
                    metadata['party_2'] = matches[0][1].strip()
                else:
                    metadata['party'] = matches[0].strip()
                break
        
        return metadata


def main():
    """Example usage of ContractParser."""
    parser = ContractParser()
    
    # Example file path - update as needed
    sample_file = Path("sample_contract.pdf")
    
    if sample_file.exists():
        try:
            contract = parser.parse_document(sample_file)
            print(f"Parsed contract: {contract.title}")
            print(f"Total clauses: {contract.total_clauses}")
            print(f"Content length: {contract.total_content_length}")
            
            for clause in contract.clauses[:3]:  # Show first 3 clauses
                print(f"\\n{clause.number}. {clause.title}")
                print(f"Content preview: {clause.content[:100]}...")
                
        except Exception as e:
            print(f"Error parsing contract: {e}")
    else:
        print("Sample contract file not found")


if __name__ == "__main__":
    main()
